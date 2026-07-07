from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.http import HttpResponse
from django.contrib import messages
from django.utils import timezone
from django.db.models import Q, Count
from datetime import datetime
import io
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from gestion_espacios_academicos.models import Solicitud
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.platypus import PageBreak
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from gestion_espacios_academicos.utils import fecha_formateada, ahora_local

@login_required
def generar_reportes(request):
    if request.method == 'GET' and any(param in request.GET for param in ['formato', 'estado', 'fecha_inicio', 'fecha_fin', 'tipo_reporte']):
        formato = request.GET.get('formato', 'pdf')
        estado = request.GET.get('estado', 'todos')
        fecha_inicio = request.GET.get('fecha_inicio')
        fecha_fin = request.GET.get('fecha_fin')
        tipo_reporte = request.GET.get('tipo_reporte', 'solicitudes')
        user = request.user

        # Solo encargados pueden generar reportes
        if user.tipo_usuario != 'encargado':
            messages.error(request, 'No tienes permiso para generar reportes.')
            return render(request, 'reportes/generar_reportes.html')

        # IDs de espacios gestionados por el usuario
        ids_carrera = user.espacios_encargados.values_list('id', flat=True)
        ids_campus = user.espacios_campus_encargados.values_list('id', flat=True)

        # Filtrar solicitudes
        solicitudes = (
            Solicitud.objects
            .select_related('usuario_solicitante', 'espacio', 'espacio_campus')
            .filter(
                Q(espacio_id__in=ids_carrera) |
                Q(espacio_campus_id__in=ids_campus)
            )
        )

        # Aplicar filtros
        if estado != 'todos':
            solicitudes = solicitudes.filter(estado=estado)
        if fecha_inicio:
            try:
                fecha_inicio_dt = datetime.strptime(fecha_inicio, '%Y-%m-%d')
                solicitudes = solicitudes.filter(fecha_evento__gte=fecha_inicio_dt)
            except ValueError:
                pass
        if fecha_fin:
            try:
                fecha_fin_dt = datetime.strptime(fecha_fin, '%Y-%m-%d')
                solicitudes = solicitudes.filter(fecha_evento__lte=fecha_fin_dt)
            except ValueError:
                pass

        # Preparar datos según el tipo de reporte
        if tipo_reporte == 'solicitudes':
            data = preparar_datos_solicitudes(solicitudes)
        elif tipo_reporte == 'ocupacion':
            data = preparar_datos_ocupacion(solicitudes)
        elif tipo_reporte == 'facultades':
            data = preparar_datos_facultades(solicitudes)
        else:
            messages.error(request, 'Tipo de reporte no válido.')
            return render(request, 'reportes/generar_reportes.html')

        # Generar archivo según formato
        if formato == 'pdf':
            return generar_pdf(data, estado, fecha_inicio, fecha_fin, tipo_reporte)
        elif formato == 'excel':
            return generar_excel_openpyxl(data, estado, fecha_inicio, fecha_fin, tipo_reporte)
        elif formato == 'word':
            return generar_word(data, estado, fecha_inicio, fecha_fin, tipo_reporte)
        else:
            messages.error(request, 'Formato no válido.')
            return render(request, 'reportes/generar_reportes.html')

    return render(request, 'reportes/generar_reportes.html')

def preparar_datos_solicitudes(solicitudes):
    """Preparar datos para el reporte de solicitudes con timezone correcto"""
    data = []
    for s in solicitudes:
        espacio_nombre = s.get_nombre_espacio()
        
        # 🔧 Usar la función helper para formatear la fecha correctamente
        fecha_str = fecha_formateada(s.fecha_evento, '%d/%m/%Y %H:%M')
        
        data.append([
            s.id,
            s.nombre_evento,
            fecha_str,  # 👈 Ahora con timezone correcto
            s.usuario_solicitante.get_full_name() or s.usuario_solicitante.username,
            espacio_nombre,
            s.estado.title(),
            s.motivo_rechazo or 'N/A'
        ])
    return data

def preparar_datos_ocupacion(solicitudes):
    """Preparar datos para el reporte de ocupación"""
    data = solicitudes.values('espacio__nombre', 'espacio_campus__nombre').annotate(
        total=Count('id'),
        aprobadas=Count('id', filter=Q(estado='aprobada'))
    ).order_by('-total')

    formatted_data = [
        [d['espacio__nombre'] or d['espacio_campus__nombre'], d['total'], d['aprobadas'], f"{d['aprobadas']/d['total']*100:.2f}%"]
        for d in data
    ]
    return formatted_data

def preparar_datos_facultades(solicitudes):
    """Preparar datos para el reporte de facultades"""
    data = solicitudes.values(
        'usuario_solicitante__facultad__nombre',
        'usuario_solicitante__carrera__nombre'
    ).annotate(
        total=Count('id'),
        aprobadas=Count('id', filter=Q(estado='aprobada'))
    ).order_by('-total')

    formatted_data = [
        [d['usuario_solicitante__facultad__nombre'], d['usuario_solicitante__carrera__nombre'], d['total'], d['aprobadas']]
        for d in data
    ]
    return formatted_data

def generar_pdf(data, estado, fecha_inicio, fecha_fin, tipo_reporte):
    """Generar reporte en formato PDF sin logo"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=landscape(A4),
        leftMargin=0.5*inch,
        rightMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    elements = []

    # Estilos
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1e3a8a'),
        alignment=TA_CENTER,
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#64748b'),
        alignment=TA_CENTER,
        spaceAfter=20
    )
    
    section_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#1e3a8a'),
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )

    # Encabezado
    elements.append(Paragraph('<b>UNIVERSIDAD AUTÓNOMA "JOSÉ BALLIVIÁN"</b>', title_style))
    elements.append(Paragraph('Sistema de Gestión de Espacios Académicos', subtitle_style))
    
    elements.append(Spacer(1, 0.2*inch))
    
    # Línea decorativa
    line_table = Table([['']], colWidths=[7.5*inch])
    line_table.setStyle(TableStyle([
        ('LINEABOVE', (0, 0), (-1, 0), 2, colors.HexColor('#fbbf24')),
    ]))
    elements.append(line_table)
    elements.append(Spacer(1, 0.2*inch))

    # Título del reporte
    title_text = {
        'solicitudes': 'REPORTE DE SOLICITUDES',
        'ocupacion': 'REPORTE DE OCUPACIÓN DE ESPACIOS',
        'facultades': 'REPORTE DE USO POR FACULTADES'
    }
    elements.append(Paragraph(title_text[tipo_reporte], title_style))

    # Información de filtros
    filter_info = f"<b>Estado:</b> {estado.title()}"
    if fecha_inicio:
        filter_info += f" | <b>Desde:</b> {fecha_inicio}"
    if fecha_fin:
        filter_info += f" | <b>Hasta:</b> {fecha_fin}"
    filter_info += f" | <b>Generado:</b> {fecha_formateada(ahora_local())}"
    
    filter_para = Paragraph(filter_info, subtitle_style)
    filter_table = Table([[filter_para]], colWidths=[7.5*inch])
    filter_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(filter_table)
    elements.append(Spacer(1, 0.3*inch))

    # Tabla de datos principales
    elements.append(Paragraph("📋 DATOS DETALLADOS", section_style))
    elements.append(Spacer(1, 0.1*inch))
    
    headers = {
        'solicitudes': ['ID', 'Evento', 'Fecha/Hora', 'Usuario', 'Espacio', 'Estado', 'Motivo Rechazo'],
        'ocupacion': ['Espacio', 'Total Solicitudes', 'Aprobadas', 'Tasa Aprobación'],
        'facultades': ['Facultad', 'Carrera', 'Total Solicitudes', 'Aprobadas']
    }
    
    table_data = [headers[tipo_reporte]]
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=7,
        leading=9,
        alignment=TA_CENTER,
        wordWrap='CJK'
    )
    
    for row in data:
        new_row = []
        for cell in row:
            new_row.append(Paragraph(str(cell), cell_style))
        table_data.append(new_row)
    
    if tipo_reporte == 'solicitudes':
        table = Table(table_data, colWidths=[25, 95, 85, 125, 105, 65, 95], rowHeights=None)
    elif tipo_reporte == 'ocupacion':
        table = Table(table_data, colWidths=[200, 120, 120, 120], rowHeights=None)
    elif tipo_reporte == 'facultades':
        table = Table(table_data, colWidths=[150, 150, 120, 120], rowHeights=None)
    else:
        table = Table(table_data, rowHeights=None)
    
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e3a8a')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')])
    ]))
    elements.append(table)
    
    # Pie de página
    elements.append(Spacer(1, 0.3*inch))
    footer_line = Table([['']], colWidths=[7.5*inch])
    footer_line.setStyle(TableStyle([
        ('LINEABOVE', (0, 0), (-1, 0), 1, colors.HexColor('#e2e8f0')),
    ]))
    elements.append(footer_line)
    
    footer_text = f'<font size=8 color="#64748b">Universidad Autónoma "José Ballivián" - Trinidad, Beni - Sistema de Gestión de Espacios Académicos</font>'
    elements.append(Paragraph(footer_text, ParagraphStyle('footer', parent=subtitle_style, alignment=TA_CENTER)))

    doc.build(elements)

    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    timestamp = ahora_local().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="reporte_{tipo_reporte}_{timestamp}.pdf"'
    return response

def generar_excel_openpyxl(data, estado, fecha_inicio, fecha_fin, tipo_reporte):
    """Generar reporte en formato Excel"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = tipo_reporte.capitalize()

    title_font = Font(name='Arial', size=18, bold=True, color='1e3a8a')
    subtitle_font = Font(name='Arial', size=11, color='64748b')
    header_font = Font(name='Arial', size=11, bold=True, color='FFFFFF')
    
    blue_fill = PatternFill(start_color='1e3a8a', end_color='1e3a8a', fill_type='solid')
    yellow_fill = PatternFill(start_color='fbbf24', end_color='fbbf24', fill_type='solid')
    light_fill = PatternFill(start_color='f8fafc', end_color='f8fafc', fill_type='solid')
    
    center_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
    border = Border(
        left=Side(style='thin', color='cbd5e1'),
        right=Side(style='thin', color='cbd5e1'),
        top=Side(style='thin', color='cbd5e1'),
        bottom=Side(style='thin', color='cbd5e1')
    )

    current_row = 1
    
    # Título universidad
    ws.merge_cells(f'A{current_row}:G{current_row}')
    title_cell = ws[f'A{current_row}']
    title_cell.value = 'UNIVERSIDAD AUTÓNOMA "JOSÉ BALLIVIÁN"'
    title_cell.font = title_font
    title_cell.alignment = center_alignment
    current_row += 1
    
    # Subtítulo
    ws.merge_cells(f'A{current_row}:G{current_row}')
    subtitle_cell = ws[f'A{current_row}']
    subtitle_cell.value = 'Sistema de Gestión de Espacios Académicos'
    subtitle_cell.font = subtitle_font
    subtitle_cell.alignment = center_alignment
    current_row += 1
    
    # Línea decorativa
    ws.merge_cells(f'A{current_row}:G{current_row}')
    line_cell = ws[f'A{current_row}']
    line_cell.fill = yellow_fill
    ws.row_dimensions[current_row].height = 5
    current_row += 2

    # Título del reporte
    title_text = {
        'solicitudes': 'REPORTE DE SOLICITUDES',
        'ocupacion': 'REPORTE DE OCUPACIÓN DE ESPACIOS',
        'facultades': 'REPORTE DE USO POR FACULTADES'
    }
    ws.merge_cells(f'A{current_row}:G{current_row}')
    report_title = ws[f'A{current_row}']
    report_title.value = title_text[tipo_reporte]
    report_title.font = Font(name='Arial', size=14, bold=True, color='1e3a8a')
    report_title.alignment = center_alignment
    current_row += 2

    # Filtros
    filter_info = f"Estado: {estado.title()}"
    if fecha_inicio:
        filter_info += f" | Desde: {fecha_inicio}"
    if fecha_fin:
        filter_info += f" | Hasta: {fecha_fin}"
    filter_info += f" | <b>Generado:</b> {fecha_formateada(ahora_local())}"
    
    ws.merge_cells(f'A{current_row}:G{current_row}')
    filter_cell = ws[f'A{current_row}']
    filter_cell.value = filter_info
    filter_cell.font = subtitle_font
    filter_cell.alignment = center_alignment
    filter_cell.fill = light_fill
    filter_cell.border = border
    current_row += 2

    # Tabla de datos
    ws.merge_cells(f'A{current_row}:G{current_row}')
    data_section = ws[f'A{current_row}']
    data_section.value = "📋 DATOS DETALLADOS"
    data_section.font = Font(name='Arial', size=12, bold=True, color='1e3a8a')
    data_section.alignment = left_alignment
    current_row += 1
    
    headers = {
        'solicitudes': ['ID', 'Evento', 'Fecha/Hora', 'Usuario', 'Espacio', 'Estado', 'Motivo Rechazo'],
        'ocupacion': ['Espacio', 'Total Solicitudes', 'Aprobadas', 'Tasa Aprobación'],
        'facultades': ['Facultad', 'Carrera', 'Total Solicitudes', 'Aprobadas']
    }
    
    # Encabezados
    for col_num, header in enumerate(headers[tipo_reporte], 1):
        cell = ws.cell(row=current_row, column=col_num)
        cell.value = header
        cell.font = header_font
        cell.fill = blue_fill
        cell.alignment = center_alignment
        cell.border = border
    
    current_row += 1

    # Datos
    for row_num, row_data in enumerate(data):
        for col_num, value in enumerate(row_data, 1):
            cell = ws.cell(row=current_row + row_num, column=col_num)
            cell.value = value
            cell.alignment = center_alignment
            cell.border = border
            cell.fill = PatternFill(
                start_color='FFFFFF' if row_num % 2 == 0 else 'f8fafc',
                end_color='FFFFFF' if row_num % 2 == 0 else 'f8fafc',
                fill_type='solid'
            )
        ws.row_dimensions[current_row + row_num].height = 30

    # Ajustar ancho de columnas
    column_widths = {
        'solicitudes': [8, 25, 18, 22, 20, 12, 22],
        'ocupacion': [30, 18, 15, 18],
        'facultades': [25, 25, 18, 15]
    }
    for col_num, width in enumerate(column_widths[tipo_reporte], 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col_num)].width = width

    # Pie de página
    current_row += len(data) + 2
    ws.merge_cells(f'A{current_row}:G{current_row}')
    footer = ws[f'A{current_row}']
    footer.value = 'Universidad Autónoma "José Ballivián" - Trinidad, Beni - Sistema de Gestión de Espacios Académicos'
    footer.font = Font(size=9, color='64748b', italic=True)
    footer.alignment = center_alignment

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    timestamp = ahora_local().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="reporte_{tipo_reporte}_{timestamp}.xlsx"'
    return response

def generar_word(data, estado, fecha_inicio, fecha_fin, tipo_reporte):
    """Generar reporte en formato Word"""
    doc = Document()
    
    # Configurar márgenes y orientación horizontal
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    
    # Título universidad
    title = doc.add_heading('UNIVERSIDAD AUTÓNOMA "JOSÉ BALLIVIÁN"', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    title.runs[0].font.size = Pt(18)
    
    # Subtítulo
    subtitle = doc.add_paragraph('Sistema de Gestión de Espacios Académicos')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    
    # Línea decorativa (simulada con tabla)
    line_table = doc.add_table(rows=1, cols=1)
    line_table.rows[0].height = Inches(0.05)
    line_cell = line_table.rows[0].cells[0]
    shading_elm = RGBColor(251, 191, 36)  # Color amarillo UABJB
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="fbbf24"/>'
    shading_elm = parse_xml(shading_xml)
    line_cell._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()
    
    # Título del reporte
    title_text = {
        'solicitudes': 'REPORTE DE SOLICITUDES',
        'ocupacion': 'REPORTE DE OCUPACIÓN DE ESPACIOS',
        'facultades': 'REPORTE DE USO POR FACULTADES'
    }
    report_title = doc.add_heading(title_text[tipo_reporte], level=1)
    report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_title.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    report_title.runs[0].font.size = Pt(14)
    
    # Filtros
    filter_info = f"Estado: {estado.title()}"
    if fecha_inicio:
        filter_info += f" | Desde: {fecha_inicio}"
    if fecha_fin:
        filter_info += f" | Hasta: {fecha_fin}"
    filter_info += f" | <b>Generado:</b> {fecha_formateada(ahora_local())}"
    filter_paragraph = doc.add_paragraph(filter_info)
    filter_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    filter_paragraph.runs[0].font.size = Pt(10)
    filter_paragraph.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph()
    
    # Tabla de datos principales
    data_heading = doc.add_heading('📋 DATOS DETALLADOS', level=2)
    data_heading.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    data_heading.runs[0].font.size = Pt(12)
    
    headers = {
        'solicitudes': ['ID', 'Evento', 'Fecha/Hora', 'Usuario', 'Espacio', 'Estado', 'Motivo Rechazo'],
        'ocupacion': ['Espacio', 'Total Solicitudes', 'Aprobadas', 'Tasa Aprobación'],
        'facultades': ['Facultad', 'Carrera', 'Total Solicitudes', 'Aprobadas']
    }
    
    table = doc.add_table(rows=len(data) + 1, cols=len(headers[tipo_reporte]))
    table.style = 'Light Grid Accent 1'
    
    # Configurar anchos de columna
    if tipo_reporte == 'solicitudes':
        column_widths = [Inches(0.4), Inches(1.3), Inches(1.1), Inches(1.5), Inches(1.3), Inches(0.8), Inches(1.2)]
    elif tipo_reporte == 'ocupacion':
        column_widths = [Inches(2.5), Inches(1.5), Inches(1.2), Inches(1.5)]
    elif tipo_reporte == 'facultades':
        column_widths = [Inches(2.0), Inches(2.0), Inches(1.5), Inches(1.2)]
    else:
        column_widths = None
    
    # Encabezados
    for i, header in enumerate(headers[tipo_reporte]):
        cell = table.rows[0].cells[i]
        cell.text = header
        if column_widths:
            cell.width = column_widths[i]
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 58, 138)
    
    # Datos
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(value)
            if column_widths:
                cell.width = column_widths[col_idx]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    
    # Pie de página
    doc.add_paragraph()
    footer_para = doc.add_paragraph('━' * 100)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.color.rgb = RGBColor(226, 232, 240)
    
    footer = doc.add_paragraph('Universidad Autónoma "José Ballivián" - Trinidad, Beni')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    footer.runs[0].font.italic = True
    
    footer2 = doc.add_paragraph('Sistema de Gestión de Espacios Académicos')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.runs[0].font.size = Pt(9)
    footer2.runs[0].font.color.rgb = RGBColor(100, 116, 139)
    footer2.runs[0].font.italic = True
    
    # Guardar en buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    timestamp = ahora_local().strftime('%Y%m%d_%H%M%S')
    response['Content-Disposition'] = f'attachment; filename="reporte_{tipo_reporte}_{timestamp}.docx"'
    return response